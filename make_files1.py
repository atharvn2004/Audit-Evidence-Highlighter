import pandas as pd
from docx import Document

# 1. Create a VALID Excel file
try:
    df = pd.DataFrame({'ID': [101, 102, 103],
                       'Details': ['Invoice A', 'Morgan Stanley Audit', 'Pending']})
    df.to_excel('auto_test.xlsx', index=False)
    print("✅ Success: Created 'auto_test.xlsx'")
except Exception as e:
    print(f"❌ Excel creation failed: {e}")

# 2. Create a VALID Word file
try:
    doc = Document()
    doc.add_heading('Audit Evidence', 0)
    doc.add_paragraph('This document is for the Morgan Stanley interview process.')
    doc.save('auto_test.docx')
    print("✅ Success: Created 'auto_test.docx'")
except Exception as e:
    print(f"❌ Word creation failed: {e}")