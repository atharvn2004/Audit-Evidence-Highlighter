from docx import Document

filename = "Morgan_Stanley_Report.docx"
doc = Document()

# 1. Add a Title
doc.add_heading('Audit Findings Report', 0)

# 2. Add a Paragraph with the target word
p = doc.add_paragraph()
p.add_run("This document confirms that ")
# We make "Morgan Stanley" bold to prove we can find styled text too
run = p.add_run("Morgan Stanley")
run.bold = True
p.add_run(" has passed the preliminary security check.")

# 3. Add a Paragraph without the target word
doc.add_paragraph("The audit was conducted on November 27, 2025.")

# 4. Add a list item with the target word
doc.add_paragraph('Reviewer: Morgan Audit Team', style='List Bullet')

doc.save(filename)
print(f"✅ Word Doc Created: {filename}")